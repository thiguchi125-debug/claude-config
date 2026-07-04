#!/usr/bin/env python3
"""議会だより 提出用docx生成スクリプト

650字確定版の原稿（JSON）から、編集委員会提出用のdocxを生成する。
書式は2026-07-02 R7.6月議会提出版（議会だより_草川卓也_R7.6一般質問_提出用.docx）を踏襲:
  - フォント: Yu Gothic・全段落bold
  - 氏名行: 11pt 黒「草川 卓也（会派：結）」
  - 大見出し（吹き出し文）: 15pt 緑 #1F5A3A
  - ■小見出し: 11.5pt 緑
  - 問: 10.5pt 緑 ／ 答: 10.5pt えんじ #B03A2E
  - 【その他の質問】: 10.5pt 緑ラベル＋黒の箇条書き（List Bullet）

使い方:
    python3 build_docx.py config.json

config.json の形式:
{
  "out": "/Users/.../議会だより_草川卓也_R◯.◯一般質問_提出用.docx",
  "name_line": "草川 卓也（会派：結）",
  "headline": "京都へ「乗換えなし」復活と、半導体の世界の亀山を",
  "sections": [
    {"head": "■ 草津線への直通運転で、京都へ乗換えなし・一直線に",
     "q": "京都へ乗換えなしで行ければ、…",
     "a": "関西本線と草津線の連携は、…"},
    ...
  ],
  "others": ["部活動の地域展開について", "中学校給食について", ...]
}
"""
import sys, json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

FONT = "Yu Gothic"
GREEN = RGBColor(0x1F, 0x5A, 0x3A)
MAROON = RGBColor(0xB0, 0x3A, 0x2E)
BLACK = None  # デフォルト色


def add_para(doc, text, size_pt, color=None, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size_pt)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if color is not None:
        r.font.color.rgb = color
    return p


def main():
    if len(sys.argv) != 2:
        print(__doc__)
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        cfg = json.load(f)

    doc = Document()
    add_para(doc, cfg["name_line"], 11)                      # sz22
    add_para(doc, cfg["headline"], 15, GREEN)                # sz30
    for sec in cfg["sections"]:
        head = sec["head"]
        if not head.startswith("■"):
            head = "■ " + head.lstrip("●").strip()
        add_para(doc, head, 11.5, GREEN)                     # sz23
        q = sec["q"]
        add_para(doc, q if q.startswith("問") else "問　" + q, 10.5, GREEN)
        a = sec["a"]
        add_para(doc, a if a.startswith("答") else "答　" + a, 10.5, MAROON)
    if cfg.get("others"):
        add_para(doc, "【その他の質問】", 10.5, GREEN)
        for o in cfg["others"]:
            add_para(doc, o, 10.5, style="List Bullet")

    doc.save(cfg["out"])

    # 字数カウント（タイトル＋見出し＋Q&A。吹き出し・氏名・その他はカウント外）
    count = 0
    for sec in cfg["sections"]:
        count += len(sec["head"].lstrip("■● ").strip()) + len(sec["q"]) + len(sec["a"])
    print("saved:", cfg["out"])
    print("見出し＋Q&A字数（吹き出し・その他カウント外）: 約%d字" % count)


if __name__ == "__main__":
    main()
